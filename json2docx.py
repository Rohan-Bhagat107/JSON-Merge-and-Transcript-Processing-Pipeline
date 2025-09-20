import os
import json
import shutil
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
 
# ------------------------ Step 1: Merging JSON Files ------------------------
 
def validator(input_dir):
    return os.path.isdir(input_dir)
 
def collect_json_data(json_path, text_container, word_container):
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    texts = data.get("text", [])
    words = data.get("words", [])
    text_container.append(texts)
    word_container.append(words)
 
def merge_json_files(file, my_text, my_words):
    with open(file, "r", encoding="utf-8") as f:
        merged_data = json.load(f)
 
    if isinstance(merged_data["text"], str):
        merged_data["text"] = [merged_data["text"]]
    if isinstance(merged_data["words"], str):
        merged_data["words"] = [merged_data["words"]]
 
    for para in my_text[1:]:
        if isinstance(para, str):
            merged_data["text"].append(para)
        elif isinstance(para, list):
            merged_data["text"].extend(para)
 
    for para in my_words[1:]:
        if isinstance(para, str):
            merged_data["words"].append(para)
        elif isinstance(para, list):
            merged_data["words"].extend(para)
 
    with open(file, "w", encoding="utf-8") as f:
        json.dump(merged_data, f, ensure_ascii=False, indent=2)
 
    print(f"✅ Successfully merged into: {file}")
    return file
 
# ------------------------ Step 2: JSON to SRT ------------------------
 
def format_time(input_sec):
    hrs = int(input_sec // 3600)
    minutes = int((input_sec % 3600) // 60)
    sec = int(input_sec % 60)
    millis = int((input_sec - int(input_sec)) * 1000)
    return f"{hrs:02}:{minutes:02}:{sec:02},{millis:03}"
 
def json_to_srt(json_file_path):
    input_dir = os.path.dirname(json_file_path)
    output_folder = os.path.join(input_dir, "Output srt")
    os.makedirs(output_folder, exist_ok=True)
 
    srt_filename = os.path.basename(json_file_path).replace('.json', '.srt')
    output_srt_path = os.path.join(output_folder, srt_filename)
 
    with open(json_file_path, "r", encoding="utf-8") as f:
        data = json.load(f)
 
    with open(output_srt_path, "w", encoding="utf-8") as file:
        seg_count = 0
        try:
            segments = [w for w in data.get("words", []) if w.get("type") == "word" and w.get("text", "").strip()]
            if not segments:
                raise Exception("No valid word segments found.")
 
            current_speaker = segments[0]["speaker_id"]
            buffer_text = segments[0]["text"]
            buffer_start = segments[0]["start"]
            buffer_end = segments[0]["end"]
 
            for segment in segments[1:]:
                speaker = segment["speaker_id"]
                text = segment["text"]
                if speaker == current_speaker:
                    buffer_text += " " + text
                    buffer_end = segment["end"]
                else:
                    seg_count += 1
                    start_time = format_time(buffer_start)
                    end_time = format_time(buffer_end)
                    file.write(f"{seg_count}\n{start_time} --> {end_time}\n[{current_speaker}] {buffer_text}\n\n")
 
                    current_speaker = speaker
                    buffer_text = text
                    buffer_start = segment["start"]
                    buffer_end = segment["end"]
 
            seg_count += 1
            start_time = format_time(buffer_start)
            end_time = format_time(buffer_end)
            file.write(f"{seg_count}\n{start_time} --> {end_time}\n[{current_speaker}] {buffer_text}\n\n")
 
        except Exception as e:
            print(f"❌ Error processing {os.path.basename(json_file_path)}: {e}")
        else:
            print(f"✅ SRT created: {output_srt_path}")
    return output_srt_path
 
# ------------------------ Step 3: SRT to DOCX ------------------------
 
def extract_speaker_dialogues_srt(srt_path):
    with open(srt_path, 'r', encoding='utf-8') as f:
        lines = [line.strip() for line in f if line.strip()]
 
    extracted = []
    current_speaker = None
    dialogue_buffer = []
 
    for line in lines:
        if line.isdigit():
            continue
        if "-->" in line:
            continue
        if line.startswith('[') and ']' in line:
            end_idx = line.find(']')
            speaker_tag = line[1:end_idx].strip()
            dialogue_text = line[end_idx+1:].strip()
            if speaker_tag != current_speaker:
                if current_speaker and dialogue_buffer:
                    extracted.append((current_speaker, ' '.join(dialogue_buffer)))
                current_speaker = speaker_tag
                dialogue_buffer = [dialogue_text] if dialogue_text else []
            else:
                dialogue_buffer.append(dialogue_text)
        else:
            if current_speaker:
                dialogue_buffer.append(line)
 
    if current_speaker and dialogue_buffer:
        extracted.append((current_speaker, ' '.join(dialogue_buffer)))
 
    return extracted
 
def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblPr.remove(tblBorders)
 
def set_fixed_table_layout(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
 
def set_cell_width(cell, width_in_inches):
    cell.width = Inches(width_in_inches)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_in_inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)
 
def is_tamil(text):
    return any('\u0B80' <= char <= '\u0BFF' for char in text)
 
def write_structured_doc(extracted_data, output_path):
    doc = Document()
 
    if not extracted_data:
        doc.add_paragraph("No speaker-tagged dialogue found.")
        doc.save(output_path)
        return
 
    for speaker, dialogue in extracted_data:
        table = doc.add_table(rows=1, cols=2)
        set_fixed_table_layout(table)
        remove_table_borders(table)
        table.autofit = False
 
        speaker_cell = table.cell(0, 0)
        dialogue_cell = table.cell(0, 1)
 
        set_cell_width(speaker_cell, 1.5)
        set_cell_width(dialogue_cell, 5.5)
 
        speaker_para = speaker_cell.paragraphs[0]
        speaker_run = speaker_para.add_run(speaker)
        speaker_run.bold = True
        speaker_run.font.size = Pt(11)
 
        dialogue_para = dialogue_cell.paragraphs[0]
        run = dialogue_para.add_run(dialogue)
        run.font.size = Pt(11)
        if is_tamil(dialogue):
            run.font.name = 'Arial Unicode MS'
 
    doc.save(output_path)
 
# ------------------------ Main Workflow ------------------------
 
if __name__ == "__main__":
    try:
        input_folder = input("📁 Enter your directory containing JSON files: ").strip()
        if not validator(input_folder):
            raise Exception()
    except Exception:
        print("❌ Please provide a valid directory path!")
    else:
        json_count = 0
        json_text = []
        json_words = []
        copied_file = None
 
        for each_file in os.listdir(input_folder):
            if each_file.lower().endswith(".json"):
                json_count += 1
                file_path = os.path.join(input_folder, each_file)
                if json_count == 1:
                    output_file = os.path.join(input_folder, f"Merged_{each_file}")
                    copied_file = shutil.copy(file_path, output_file)
                    collect_json_data(file_path, json_text, json_words)
                else:
                    collect_json_data(file_path, json_text, json_words)
 
        if copied_file:
            merged_path = merge_json_files(copied_file, json_text, json_words)
            srt_path = json_to_srt(merged_path)
 
            output_docx_folder = os.path.join(input_folder, "Srt to docx output")
            os.makedirs(output_docx_folder, exist_ok=True)
            docx_filename = os.path.basename(srt_path).replace(".srt", ".docx")
            docx_path = os.path.join(output_docx_folder, docx_filename)
 
            extracted = extract_speaker_dialogues_srt(srt_path)
            write_structured_doc(extracted, docx_path)
 
            print(f"📄 Final DOCX saved: {docx_path}")
        else:
            print("⚠️ No JSON files found to merge.")
 
    finally:
        print("✅ Done. Thank you!")
